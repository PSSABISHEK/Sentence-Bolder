import tkinter as tk
import os.path
from docx import Document
from docx.shared import Inches

def show_entry_fields():
    if os.path.isfile('./InputFolder/'+e1.get()+'.docx'):
        doc = Document('./InputFolder/'+e1.get()+'.docx')
        doc1 = Document()

        for i in doc.paragraphs:
            if 'M:' in i.text:
                p = doc1.add_paragraph()
                p.add_run(i.text).bold = True
            else:
                p = doc1.add_paragraph()
                p.add_run(i.text)

        doc1.save('./OutputFolder/'+e2.get()+'.docx')
        if os.path.isfile('./OutputFolder/'+e2.get()+'.docx'):
            tk.Label(master, text= '                                                      ', fg='green', font=('helvetica', 12, 'bold')).grid(row=2)
            tk.Label(master, text= 'File Created', fg='green', font=('helvetica', 12, 'bold')).grid(row=2)
        else:
            tk.Label(master, text= '                                           ', fg='green', font=('helvetica', 12, 'bold')).grid(row=2)
            tk.Label(master, text= 'Error Occured', fg='green', font=('helvetica', 12, 'bold')).grid(row=2) 
    else:
        tk.Label(master, text= '                                                    ', fg='green', font=('helvetica', 12, 'bold')).grid(row=2)
        tk.Label(master, text= 'Input file does not exist', fg='green', font=('helvetica', 12, 'bold')).grid(row=2)

master = tk.Tk()
tk.Label(master, 
         text="Input File Name").grid(row=0)
tk.Label(master, 
         text="Output File Name").grid(row=1)

e1 = tk.Entry(master)
e2 = tk.Entry(master)

e1.grid(row=0, column=1)
e2.grid(row=1, column=1)

tk.Button(master, 
          text='Quit', 
          command=master.quit).grid(row=3,
                                    column=0, 
                                    sticky=tk.W, 
                                    pady=4)
tk.Button(master, 
          text='Submit', command=show_entry_fields).grid(row=3, 
                                                       column=1, 
                                                       sticky=tk.W, 
                                                       pady=4)

tk.mainloop()